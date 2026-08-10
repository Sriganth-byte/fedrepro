from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "TECHNICAL_IMPLEMENTATION_REPORT.md"
OUTPUT = ROOT / "docs" / "FedRepro_Phase_1_Implementation_Report.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "687386"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"
BORDER = "B9C4D0"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_bottom_border(paragraph, color=BLUE, size=12):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    tokens = {
        "Title": (27, DARK_BLUE, 0, 7),
        "Subtitle": (13, MUTED, 0, 12),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    code = styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abstract = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    next_abstract = max(existing_abstract, default=0) + 1
    next_num = max(existing_num, default=0) + 1

    def create(kind, abstract_id, num_id):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        level.extend([start, fmt, text, justification, p_pr])
        abstract.append(level)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    create("bullet", next_abstract, next_num)
    create("decimal", next_abstract + 1, next_num + 1)
    return next_num, next_num + 1


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_inline(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, "Consolas", 9.5, DARK_BLUE)
        else:
            paragraph.add_run(part)


def table_widths(rows):
    cols = len(rows[0])
    if cols == 2:
        left_lengths = [len(row[0]) for row in rows]
        if max(left_lengths) < 18:
            return [2700, 6660]
        return [3900, 5460]
    if cols == 3:
        return [1800, 3000, 4560]
    if cols == 4:
        return [1500, 3100, 1700, 3060]
    base = 9360 // cols
    return [base] * (cols - 1) + [9360 - base * (cols - 1)]


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = table_widths(rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for row_index, source_row in enumerate(rows):
        for col_index, value in enumerate(source_row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline(paragraph, value)
            for run in paragraph.runs:
                set_font(run, "Calibri", 9.4, INK, bold=row_index == 0)
            if row_index == 0:
                shade_cell(cell, PALE_BLUE)
            elif row_index % 2 == 0:
                shade_cell(cell, "FAFBFC")
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def configure_page(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("FEDREPRO  |  PHASE 1 IMPLEMENTATION REPORT")
    set_font(run, "Calibri", 8.5, MUTED, bold=True)
    add_bottom_border(p, BORDER, 5)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(3)
    run = p.add_run("FedRepro 1.0.0  •  5 July 2026  •  Page ")
    set_font(run, "Calibri", 8.5, MUTED)
    add_page_field(p)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("TECHNICAL ASSESSMENT")
    set_font(run, "Calibri", 10, BLUE, bold=True)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_before = Pt(8)
    title.add_run("FedRepro Phase 1\nImplementation Report")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Architecture, evidence lifecycle, deterministic analysis, verification, and production-readiness assessment")

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    add_bottom_border(rule, BLUE, 16)

    metadata = [
        ("Assessment date", "5 July 2026"),
        ("System", "FedRepro 1.0.0"),
        ("Implementation scope", "Backend, database, API, frontend, security, verification"),
        ("Assessment basis", "Repository inspection and executable checks"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        label_run = p.add_run(f"{label}: ")
        set_font(label_run, "Calibri", 10, DARK_BLUE, bold=True)
        value_run = p.add_run(value)
        set_font(value_run, "Calibri", 10, INK)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT)
    p_pr.append(shading)
    lead = p.add_run("Assessment outcome. ")
    set_font(lead, "Calibri", 11, DARK_BLUE, bold=True)
    rest = p.add_run(
        "Functionally complete for Phase 1 research use, with a sound evidence-first design. "
        "Production deployment requires security, operations, scalability, and test-coverage hardening."
    )
    set_font(rest, "Calibri", 11, INK)
    doc.add_page_break()


def strip_front_matter(lines):
    first_heading_seen = False
    result = []
    for line in lines:
        if line.startswith("# "):
            if not first_heading_seen:
                first_heading_seen = True
                continue
        if first_heading_seen and (
            line.startswith("**Assessment date:**")
            or line.startswith("**System:**")
            or line.startswith("**Scope:**")
        ):
            continue
        result.append(line)
    return result


def build():
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    bullet_num, decimal_num = add_numbering(doc)
    add_header_footer(doc.sections[0])
    add_cover(doc)

    lines = strip_front_matter(SOURCE.read_text(encoding="utf-8").splitlines())
    index = 0
    code_mode = False
    code_lines = []
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if code_mode:
                p = doc.add_paragraph(style="Code Block")
                shade = OxmlElement("w:shd")
                shade.set(qn("w:fill"), LIGHT)
                p._p.get_or_add_pPr().append(shade)
                p.add_run("\n".join(code_lines))
                code_lines = []
                code_mode = False
            else:
                code_mode = True
            index += 1
            continue
        if code_mode:
            code_lines.append(line)
            index += 1
            continue
        if not line.strip():
            index += 1
            continue
        if line.startswith("|"):
            rows = []
            while index < len(lines) and lines[index].startswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", cell) for cell in cells):
                    rows.append(cells)
                index += 1
            if rows:
                add_table(doc, rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1)) - 1
            p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
            add_inline(p, heading.group(2))
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            apply_num(p, bullet_num)
            add_inline(p, bullet.group(1))
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            apply_num(p, decimal_num)
            add_inline(p, numbered.group(1))
            index += 1
            continue
        p = doc.add_paragraph()
        add_inline(p, line)
        index += 1

    configure_page(doc)
    for section in doc.sections:
        add_header_footer(section)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "FedRepro Phase 1 Implementation Report"
    doc.core_properties.subject = "Technical implementation and production-readiness assessment"
    doc.core_properties.author = "FedRepro Engineering Assessment"
    doc.core_properties.keywords = "FedRepro, implementation, architecture, reproducibility, ML research"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
