from __future__ import annotations

import json
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.orm import Session

from app.api.routes.ai import resolve_evidence
from app.models.entities import Dataset, DatasetConfiguration, DatasetRegistration, DatasetVersion, DiagnosisReport, SemanticDiffReport, Study
from app.services.ai_explanation_service import AIExplanationService


class ExecutiveReportService:
    def __init__(self, db: Session):
        self.db = db

    def build(self, study: Study, include_ai: bool = False) -> Path:
        datasets = self.db.query(Dataset).filter(Dataset.study_id == study.id).all()
        versions = (
            self.db.query(DatasetVersion)
            .join(Dataset)
            .filter(Dataset.study_id == study.id)
            .order_by(DatasetVersion.created_at)
            .all()
        )
        latest = versions[-1] if versions else None
        doc = Document()
        self._styles(doc)
        self._title(doc, study)
        self._summary(doc, study, datasets, versions)
        self._protocol(doc, study)
        self._pipeline(doc, datasets, versions)
        self._evidence(doc, datasets)
        self._versions(doc, versions)
        self._recreation(doc, latest)
        self._metrics(doc, versions)
        self._inferences(doc, versions)
        self._next_actions(doc, datasets, versions)
        if include_ai and latest:
            self._ai_interpretation(doc, study, latest)
        output = Path(tempfile.gettempdir()) / f"fedrepro-study-{study.id}-executive-report.docx"
        doc.save(output)
        return output

    def _styles(self, doc: Document) -> None:
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        for name, size, color in (("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")):
            style = doc.styles[name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)

    def _title(self, doc: Document, study: Study) -> None:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("FedRepro Executive Study Report")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor.from_string("0B2545")
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(f"{study.name} | {study.ml_task} | Study #{study.id}")
        doc.add_paragraph()

    def _summary(self, doc: Document, study: Study, datasets: list[Dataset], versions: list[DatasetVersion]) -> None:
        doc.add_heading("Executive Summary", level=1)
        latest = versions[-1] if versions else None
        text = (
            f"This report summarizes the current state of the FedRepro study '{study.name}'. "
            f"The workspace currently contains {len(datasets)} dataset record(s) and {len(versions)} immutable dataset version(s)."
        )
        if latest:
            text += f" The latest version is V{latest.version_number} with {latest.row_count} rows and {latest.column_count} columns."
        else:
            text += " No immutable dataset version has been configured yet."
        doc.add_paragraph(text)

    def _protocol(self, doc: Document, study: Study) -> None:
        doc.add_heading("Study Protocol", level=1)
        self._kv_table(doc, [
            ("Study name", study.name),
            ("ML task", study.ml_task),
            ("Description", study.description or "Not documented"),
            ("Research objective", study.problem_objective or "Not documented"),
            ("Intended use", study.intended_use_case or "Not documented"),
        ])

    def _pipeline(self, doc: Document, datasets: list[Dataset], versions: list[DatasetVersion]) -> None:
        doc.add_heading("Pipeline Status", level=1)
        rows = [
            ("Study protocol", "Complete"),
            ("Dataset evidence registered", "Complete" if datasets else "Not started"),
            ("Immutable version created", "Complete" if versions else "Not started"),
            ("Fingerprint generated", "Complete" if any(getattr(v, "fingerprint", None) for v in versions) else "Not available"),
            ("Diagnosis generated", "Complete" if any(self._diagnosis(v.id) for v in versions) else "Not available"),
        ]
        self._table(doc, ("Step", "Status"), rows)

    def _evidence(self, doc: Document, datasets: list[Dataset]) -> None:
        doc.add_heading("Dataset Evidence Registry", level=1)
        rows = []
        for dataset in datasets:
            registrations = self.db.query(DatasetRegistration).filter(DatasetRegistration.dataset_id == dataset.id).all()
            rows.append((dataset.name, str(len(registrations)), str(len(dataset.versions)), str(dataset.created_at.date())))
        if rows:
            self._table(doc, ("Dataset", "Registrations", "Versions", "Created"), rows)
        else:
            doc.add_paragraph("No dataset evidence has been registered.")

    def _versions(self, doc: Document, versions: list[DatasetVersion]) -> None:
        doc.add_heading("Version Ledger and Fingerprints", level=1)
        rows = []
        for version in versions:
            fingerprint = version.fingerprint.combined_fingerprint if version.fingerprint else "Not available"
            rows.append((f"V{version.version_number}", str(version.parent_version_id or "Baseline"), str(version.row_count), str(version.column_count), version.file_hash[:16], fingerprint[:16]))
        if rows:
            self._table(doc, ("Version", "Parent", "Rows", "Columns", "File hash", "Fingerprint"), rows)
        else:
            doc.add_paragraph("No immutable versions have been created.")

    def _recreation(self, doc: Document, latest: DatasetVersion | None) -> None:
        doc.add_heading("Latest Version Recreation Evidence", level=1)
        if not latest:
            doc.add_paragraph("No immutable version is available for recreation evidence.")
            return
        configuration = self.db.get(DatasetConfiguration, latest.configuration_id)
        fingerprint = latest.fingerprint
        required = [
            ("Source file bytes", latest.file_hash),
            ("Combined fingerprint", fingerprint.combined_fingerprint if fingerprint else None),
            ("Configuration hash", configuration.configuration_hash if configuration else None),
            ("Fingerprint ruleset", fingerprint.algorithm_version if fingerprint else None),
            ("Target column", configuration.target_column if configuration else None),
            ("Primary metric", configuration.primary_metric if configuration else None),
            ("Validation strategy", configuration.validation_strategy if configuration else None),
            ("Feature selection", configuration.feature_selection_mode if configuration else None),
            ("Scaling strategy", configuration.scaling_strategy if configuration else None),
        ]
        rows = [(label, str(value)) for label, value in required if value]
        if rows:
            self._table(doc, ("Recreation input", "Recorded value"), rows)
        else:
            doc.add_paragraph(f"V{latest.version_number} does not yet expose recreation evidence.")

    def _metrics(self, doc: Document, versions: list[DatasetVersion]) -> None:
        doc.add_heading("Available Metrics", level=1)
        rows = []
        for version in versions:
            diagnosis = self._diagnosis(version.id)
            semantic = self.db.query(SemanticDiffReport).filter(SemanticDiffReport.current_version_id == version.id).first()
            rows.append((
                f"V{version.version_number}",
                self._metric(diagnosis.mlrs_score if diagnosis else None),
                self._metric(diagnosis.lrs_score if diagnosis else None),
                self._metric(semantic.scm_score if semantic else None),
                self._metric(semantic.dsi_score if semantic else None),
                str(len(diagnosis.findings_json)) if diagnosis else "Not available",
            ))
        if rows:
            self._table(doc, ("Version", "MLRS", "LRS", "SCM", "DSI", "Findings"), rows)
        else:
            doc.add_paragraph("No metrics are available yet because no version has been analyzed.")

    def _inferences(self, doc: Document, versions: list[DatasetVersion]) -> None:
        doc.add_heading("Evidence-Bound Inferences", level=1)
        if not versions:
            doc.add_paragraph("Dataset evidence has not yet reached analysis, so no data-quality inferences are available.")
            return
        for version in versions:
            diagnosis = self._diagnosis(version.id)
            if not diagnosis:
                doc.add_paragraph(f"V{version.version_number}: No diagnosis report is available.")
                continue
            findings = diagnosis.findings_json or []
            if not findings:
                doc.add_paragraph(f"V{version.version_number}: No material diagnosis findings crossed deterministic thresholds.")
            for finding in findings[:5]:
                doc.add_paragraph(f"V{version.version_number}: {finding.get('issue', 'Finding')} - {finding.get('risk', 'No risk text recorded')}", style="List Bullet")

    def _next_actions(self, doc: Document, datasets: list[Dataset], versions: list[DatasetVersion]) -> None:
        doc.add_heading("Recommended Next Actions", level=1)
        actions = []
        if not datasets:
            actions.append("Register baseline dataset evidence.")
        elif not versions:
            actions.append("Configure target, metric, validation, and feature settings to create an immutable version.")
        else:
            actions.append("Review diagnosis and semantic changes for the latest version.")
            actions.append("Generate or register controlled dataset variants for comparison.")
            actions.append("Use exported reports to document reproducibility evidence.")
        for action in actions:
            doc.add_paragraph(action, style="List Bullet")

    def _ai_interpretation(self, doc: Document, study: Study, latest: DatasetVersion) -> None:
        doc.add_heading("Optional AI Interpretation", level=1)
        try:
            evidence = resolve_evidence(self.db, study, "version_analysis", latest.id)
            result = AIExplanationService(self.db).explain(study, "version_analysis", latest.id, evidence)
            content = json.loads(result.content)
            self._ai_section(doc, "Executive summary", content.get("executive_summary"))
            self._ai_list(doc, "Selected version profile", content.get("selected_version_profile", []))
            for item in content.get("version_evolution", []):
                doc.add_heading(item.get("transition", "Version transition"), level=3)
                self._bullets(doc, item.get("changes", []))
                if item.get("interpretation"):
                    doc.add_paragraph(f"Interpretation: {item['interpretation']}")
            self._ai_list(doc, "Research cautions", content.get("research_cautions", []))
            self._ai_list(doc, "Reproducibility recipe", content.get("reproducibility_recipe", []))
            self._ai_list(doc, "Potential effects", content.get("potential_effects", []))
            self._ai_list(doc, "Recommended next checks", content.get("recommended_next_checks", []))
            self._ai_section(doc, "Conclusion", content.get("conclusion"))
        except Exception as exc:
            doc.add_paragraph(f"AI interpretation was not included: {exc}")

    def _diagnosis(self, version_id: int) -> DiagnosisReport | None:
        return self.db.query(DiagnosisReport).filter(DiagnosisReport.version_id == version_id).first()

    @staticmethod
    def _metric(value) -> str:
        return "Not available" if value is None else f"{float(value):.2f}"

    def _kv_table(self, doc: Document, rows: list[tuple[str, str]]) -> None:
        self._table(doc, ("Field", "Value"), rows)

    def _ai_section(self, doc: Document, title: str, text: str | None) -> None:
        if not text:
            return
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)

    def _ai_list(self, doc: Document, title: str, rows: list[str]) -> None:
        if not rows:
            return
        doc.add_heading(title, level=2)
        self._bullets(doc, rows)

    @staticmethod
    def _bullets(doc: Document, rows: list[str]) -> None:
        for row in rows:
            if row:
                doc.add_paragraph(str(row), style="List Bullet")

    @staticmethod
    def _table(doc: Document, headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
        doc.add_paragraph()
