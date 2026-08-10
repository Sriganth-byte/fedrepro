from __future__ import annotations

import json
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.orm import Session

from app.models.entities import DatasetProfileReport, DiagnosisReport
from app.services.diagnosis_contract_service import DiagnosisContractService


class DiagnosisReportExportService:
    def __init__(self, db: Session):
        self.db = db

    def build(self, study, version) -> Path:
        diagnosis = self.db.query(DiagnosisReport).filter(DiagnosisReport.version_id == version.id).first()
        profile = self.db.query(DatasetProfileReport).filter(DatasetProfileReport.version_id == version.id).first()
        if not diagnosis:
            raise ValueError("Diagnosis report not found")
        contract = DiagnosisContractService().build(study, version, profile, diagnosis)
        doc = Document()
        self._styles(doc)
        self._title(doc, contract)
        self._summary(doc, contract)
        self._risk_families(doc, contract)
        self._interventions(doc, contract)
        self._decisions(doc, contract)
        self._plan(doc, contract)
        self._handoff(doc, contract)
        self._findings(doc, contract)
        self._columns(doc, contract)
        self._appendix(doc, contract)
        output = Path(tempfile.gettempdir()) / f"fedrepro-diagnosis-v{version.id}-report.docx"
        doc.save(output)
        return output

    def _styles(self, doc: Document) -> None:
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.1
        for name, size, color in (("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")):
            style = doc.styles[name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)

    def _title(self, doc: Document, contract: dict) -> None:
        header = contract["header"]
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("FedRepro Diagnosis and Intervention Report")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor.from_string("0B2545")
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(f"{header['study_name']} | {header['dataset_name']} | V{header['version_number']}")
        doc.add_paragraph()

    def _summary(self, doc: Document, contract: dict) -> None:
        header, readiness = contract["header"], contract["readiness"]
        doc.add_heading("1. Diagnosis Header", level=1)
        self._table(doc, ("Field", "Value"), [
            ("Study", header["study_name"]),
            ("Task", header["ml_task"]),
            ("Dataset", header["dataset_name"]),
            ("Version", f"V{header['version_number']}"),
            ("Diagnosis report", header["diagnosis_report_id"]),
            ("Profile report", header.get("profile_report_id") or "Not available"),
            ("Fingerprint", self._short(header.get("dataset_fingerprint"))),
            ("Diagnosis ruleset", header["diagnosis_ruleset"]),
        ])
        doc.add_heading("2. Diagnosis Readiness Summary", level=1)
        self._table(doc, ("Signal", "Value"), [
            ("Status", readiness["status"]),
            ("MLRS", readiness["mlrs_score"]),
            ("LRS", readiness["lrs_score"]),
            ("Findings", readiness["finding_count"]),
            ("Risk families", ", ".join(readiness["detected_risk_families"])),
            ("Intervention options", readiness["intervention_count"]),
            ("Required decisions", readiness["required_decision_count"]),
        ])

    def _risk_families(self, doc: Document, contract: dict) -> None:
        doc.add_heading("3. Risk Family Overview", level=1)
        rows = [(item["family"], item["severity"], item["finding_count"], len(item["affected_columns"]), "Yes" if item["intervention_available"] else "No") for item in contract["risk_families"]]
        self._table_or_empty(doc, ("Family", "Severity", "Findings", "Columns", "Intervention"), rows, "No diagnosis risk families were detected.")

    def _interventions(self, doc: Document, contract: dict) -> None:
        doc.add_heading("4. Dynamic Intervention Options", level=1)
        if not contract["intervention_options"]:
            doc.add_paragraph("No intervention options were generated because no actionable diagnosis findings were available.")
            return
        for option in contract["intervention_options"]:
            doc.add_heading(option["title"], level=2)
            doc.add_paragraph(option["objective"])
            self._table(doc, ("Field", "Value"), [
                ("Triggered by", ", ".join(option["source_findings"])),
                ("Affected columns", ", ".join(option["affected_columns"]) or "Dataset-level"),
                ("Safety/status", option["status"]),
                ("Recommended comparison", option["recommended_comparison"]),
            ])
            self._table(doc, ("Operation", "Purpose", "Columns"), [(op["operation"], op["purpose"], ", ".join(op["columns"]) or "Dataset-level") for op in option["operations"]])
            self._bullets(doc, "Expected dataset changes", option["expected_changes"])
            self._metric_impact(doc, option["metric_impact"])
            self._bullets(doc, "Risks introduced", option["risks_introduced"])

    def _metric_impact(self, doc: Document, impact: dict) -> None:
        doc.add_heading("Metric Impact Preview", level=3)
        self._table(doc, ("Impact area", "Evidence-bound preview"), [
            ("Affected metrics", ", ".join(impact.get("affected_metrics", []))),
            ("Possible positive effect", impact.get("possible_positive_effect", "")),
            ("Possible negative effect", impact.get("possible_negative_effect", "")),
            ("Reliability effect", impact.get("reliability_effect", "")),
            ("Final finding implication", impact.get("final_finding_implication", "")),
        ])
        self._bullets(doc, "Verification required", impact.get("verification_required", []))

    def _decisions(self, doc: Document, contract: dict) -> None:
        doc.add_heading("5. Human Decision Queue", level=1)
        rows = [(item["question"], ", ".join(item["affected_columns"]) or "Dataset-level", item["recommended_default"]) for item in contract["human_decisions"]]
        self._table_or_empty(doc, ("Decision", "Scope", "Recommended default"), rows, "No user approvals are required by the current intervention plan.")

    def _plan(self, doc: Document, contract: dict) -> None:
        doc.add_heading("6. Selected Variant Plan", level=1)
        plan = contract["selected_variant_plan"]
        self._table(doc, ("Field", "Value"), [
            ("Source version", f"V{plan['source_version_number']}"),
            ("Selected options", len(plan["selected_option_ids"])),
            ("Operations", plan["operation_count"]),
            ("Affected columns", ", ".join(plan["affected_columns"]) or "Dataset-level"),
            ("Unresolved decisions", len(plan["unresolved_decisions"])),
            ("Variant names", "; ".join(plan["variant_names"]) or "No variants planned"),
        ])

    def _handoff(self, doc: Document, contract: dict) -> None:
        doc.add_heading("7. Experiment Handoff Contract", level=1)
        handoff = contract["experiment_handoff"]
        self._table(doc, ("Field", "Value"), [
            ("Source version id", handoff["source_version_id"]),
            ("Diagnosis report id", handoff["diagnosis_report_id"]),
            ("Task type", handoff["task_type"]),
            ("Required baseline", handoff["required_baseline"]),
            ("Recommended metrics", ", ".join(handoff["recommended_metrics"])),
        ])
        self._bullets(doc, "Experiment constraints", handoff["constraints"])
        self._bullets(doc, "Experiment cautions", handoff["cautions"])

    def _findings(self, doc: Document, contract: dict) -> None:
        doc.add_heading("8. Findings Board", level=1)
        rows = [(item.get("code"), item.get("severity"), item.get("issue"), item.get("risk"), item.get("recommendation")) for item in contract["findings"]]
        self._table_or_empty(doc, ("Code", "Severity", "Issue", "Risk", "Recommendation"), rows, "No diagnosis findings crossed deterministic thresholds.")

    def _columns(self, doc: Document, contract: dict) -> None:
        doc.add_heading("9. Column Impact Matrix", level=1)
        rows = [(item["column"], item.get("role"), item.get("data_type"), ", ".join(item["risk_families"]), item["recommended_operation_count"]) for item in contract["column_impact"]]
        self._table_or_empty(doc, ("Column", "Role", "Type", "Risk families", "Operations"), rows, "No column-specific diagnosis impacts were available.")

    def _appendix(self, doc: Document, contract: dict) -> None:
        doc.add_heading("Appendix: Structured Diagnosis Contract", level=1)
        doc.add_paragraph(json.dumps(contract, default=str, indent=2))

    @staticmethod
    def _short(value) -> str:
        return f"{str(value)[:16]}..." if value else "Not available"

    def _table_or_empty(self, doc: Document, headers: tuple[str, ...], rows: list[tuple], empty: str) -> None:
        if rows:
            self._table(doc, headers, rows)
        else:
            doc.add_paragraph(empty)

    def _bullets(self, doc: Document, title: str, rows: list[str]) -> None:
        if not rows:
            return
        doc.add_heading(title, level=3)
        for row in rows:
            doc.add_paragraph(str(row), style="List Bullet")

    @staticmethod
    def _table(doc: Document, headers: tuple[str, ...], rows: list[tuple]) -> None:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = str(header)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = str(value)
        doc.add_paragraph()
